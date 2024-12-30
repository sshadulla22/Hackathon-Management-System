from xml.dom.minidom import Document
import streamlit as st
import sqlite3
from datetime import datetime
import pandas as pd
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # Correct import here
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
import plotly.express as px
from datetime import datetime, timedelta
from fpdf import FPDF
from docx import Document
from reportlab.pdfgen import canvas
from sympy import content


# Page Configurations
st.set_page_config(page_title="Hackathon Management System", layout="wide")

# Database Setup
conn = sqlite3.connect('hackathon.db')
cursor = conn.cursor()

# Create Tables
cursor.execute('''
    CREATE TABLE IF NOT EXISTS organizing_team (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        role TEXT,
        team TEXT,
        contact TEXT,
        departmen TEXT
              
    )
''')

cursor.execute('''
    CREATE TABLE IF NOT EXISTS tasks (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        task_name TEXT,
        assigned_to TEXT,
        due_date DATE,
        status TEXT DEFAULT 'Pending',
        priority TEXT,
        team TEXT
               
    )
''')
conn.commit()


# import streamlit as st

# # Display Logo on App Startup
# def show_logo_interface():
#     st.image('logo.png', use_column_width=True)  # Replace with your logo file path
#     st.title("Welcome to the Organizing Team Management App")
#     st.write("This is the home page of the app where we manage the team members and tasks for the event.")

# # Main app logic
# def main():
#     menu = st.sidebar.selectbox("Select Menu", ["Home", "Organizing Team", "Tasks Management"])

#     if menu == "Home":
#         show_logo_interface()  # Display the logo interface when the app starts
#     elif menu == "Organizing Team":
#         # Your code for organizing team management
#         pass
#     elif menu == "Tasks Management":
#         # Your code for tasks management
#         pass

# if __name__ == "__main__":
#     main()


import streamlit as st

# Sidebar Title and Logo
st.sidebar.markdown(
    """
    <div style="text-align: center;">
        <h1 style="color: #0068c9;
        font-size: 45px;
        font-weight: bold;">Algorithm 9.0</h1>
        <p style="font-size: 14px; color: #555;">Your Hackathon Assistant</p>
    </div>
    """, 
    unsafe_allow_html=True
)
st.sidebar.image("Group 50.png", use_column_width=True)  # Replace with your logo

st.sidebar.markdown("---")

# Sidebar Navigation
st.sidebar.markdown("<h2 style='color: #4CAF50;'>📌 Navigation Menu</h2>", unsafe_allow_html=True)

menu = st.sidebar.radio(
    "Select an Option",
    [
        "📈 Dashboard",          # Add icon and simplify text
        "👥 Organizing Team",
        "✅ Tasks Management",
        "🛡️ Team Pages",
        "📂 Export Data",
        "📝 Attendence"        # New Menu Item Added
    ]
)

# Add Contextual Information
if menu == " 📈 Dashboard":
    st.sidebar.markdown("<p style='font-size: 14px; color: #777;'>Overview of your progress and analytics.</p>", unsafe_allow_html=True)
elif menu == "👥 Organizing Team":
    st.sidebar.markdown("<p style='font-size: 14px; color: #777;'>Manage and collaborate with your organizing team.</p>", unsafe_allow_html=True)
elif menu == "✅ Tasks Management":
    st.sidebar.markdown("<p style='font-size: 14px; color: #777;'>Track and update the progress of tasks.</p>", unsafe_allow_html=True)
elif menu == "🛡️ Team Pages":
    st.sidebar.markdown("<p style='font-size: 14px; color: #777;'>Collaborate and communicate with your team members.</p>", unsafe_allow_html=True)
elif menu == "📂 Export Data":
    st.sidebar.markdown("<p style='font-size: 14px; color: #777;'>Download reports and data files easily.</p>", unsafe_allow_html=True)
elif menu == "💬 Attendence":
    st.sidebar.markdown("<p style='font-size: 14px; color: #777;'>Get help and support for any queries you may have.</p>", unsafe_allow_html=True)

st.sidebar.markdown("---")

# Footer or Quick Links
st.sidebar.markdown(
    """
    <div style="text-align: center;">
        <h4 style="color: #2196F3;">🔗 Quick Links</h4>
        <a href="https://hackathon.com" target="_blank">Hackathon Guide</a> |
        <a href="https://streamlit.io" target="_blank">Streamlit Docs</a>
        <p style="font-size: 12px; color: #AAA;">Developed by Shadulla shaikh</p>
    </div>
    """, 
    unsafe_allow_html=True
)

# Dashboard
if menu == "📈 Dashboard":
    st.title("Hackathon Management Dashboard")

    # Layout for displaying stats in a more organized way
    col1, col2, col3 = st.columns(3)
    
    # Fetch Stats
    team_count = cursor.execute("SELECT COUNT(*) FROM organizing_team").fetchone()[0]
    pending_tasks = cursor.execute("SELECT COUNT(*) FROM tasks WHERE status='Pending'").fetchone()[0]
    completed_tasks = cursor.execute("SELECT COUNT(*) FROM tasks WHERE status='Completed'").fetchone()[0]

    # Display Stats with Metrics
    with col1:
        st.metric("Total Organizing Team Members", team_count)
    with col2:
        st.metric("Pending Tasks", pending_tasks)
    with col3:
        st.metric("Completed Tasks", completed_tasks)

    # Visualization: Pie Chart for Task Status
    task_status_data = {
        "Status": ["Pending", "Completed"],
        "Count": [pending_tasks, completed_tasks]
    }
    task_status_df = pd.DataFrame(task_status_data)
    fig = px.pie(task_status_df, names="Status", values="Count", title="Task Status Distribution")
    st.plotly_chart(fig)

    # Visualization: Bar Chart for Team Distribution
    team_data = cursor.execute("SELECT team, COUNT(*) FROM tasks GROUP BY team").fetchall()
    teams_df = pd.DataFrame(team_data, columns=["Team", "Task Count"])
    bar_fig = px.bar(teams_df, x="Team", y="Task Count", title="Tasks Distribution by Team", color="Team")
    st.plotly_chart(bar_fig)


    # # Countdown Timer
    # st.subheader("Time Remaining until Hackathon")

    # # Set hackathon end date and calculate the time remaining
    # hackathon_end_date = datetime(2024, 11, 27, 18, 0)  # Example end date: November 27, 2024, 18:00
    # current_time = datetime.now()
    # time_remaining = hackathon_end_date - current_time

    # # Display countdown
    # if time_remaining > timedelta(0):
    #     hours, remainder = divmod(time_remaining.seconds, 3600)
    #     minutes, seconds = divmod(remainder, 60)
    #     st.markdown(f"**{time_remaining.days} days, {hours:02}:{minutes:02}:{seconds:02} remaining**")
    # else:
    #     st.markdown("**Hackathon has ended!**")

    # # Progress Bar for Pending Tasks
    # st.subheader("Progress of Pending Tasks")
    # total_tasks = team_count + pending_tasks + completed_tasks  # Example calculation
    # if total_tasks > 0:
    #     progress = (completed_tasks / total_tasks) * 100
    #     st.progress(progress)
    # else:
    #     st.write("No tasks available.")







    # Custom CSS for enhanced UI
    st.markdown(
        """
        <style>
        .stButton button {
            background-color: #1e90ff;
            color: white;
            border-radius: 10px;
            padding: 10px 20px;
            font-size: 16px;
        }
        .stButton button:hover {
            background-color: #4682b4;
        }
        .stMetric {
            font-size: 22px;
            font-weight: bold;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Displaying Background Color & Styling the Page
    st.markdown(
        """
        <style>
        body {
            background-color: #f4f7fa;
        }
        .css-1v3fvcr {
            background-color: #ffffff;
            padding: 10px;
            border-radius: 8px;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        }
        .css-1dpxtcx {
            margin: 20px 0;
        }
        </style>
        """,
        unsafe_allow_html=True



        
    )


# Organizing Team Management
if menu == "👥 Organizing Team":
    st.title("Organizing Team Management")
    
    # Add Team Member Form
    with st.form("add_team_member_form"):
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            name = st.text_input("Name (Roll No.)")
        with col2:
            role = st.selectbox("Role", ["Event Head", "Guide", "Team Lead", "Member", "Volunteer","Faculty Co-ordinator"])
        with col3:
            team = st.selectbox("Team Name", [
                "Event Management Team",
                "Creative Team",
                "Design Team",
                "Marketing Team",
                "Sponsorship Team",
                "Hospitality Team",
                "Media Team",
                "Documentation Team",
                "Tech Team"
            ])
        with col4:
            contact = st.text_input("Contact (Phone/Email)")
        
        with col5:
            department = st.text_input("department")   
        
        submitted = st.form_submit_button("Add Team Member")

        if submitted and name and role:
            cursor.execute("INSERT INTO organizing_team (name, role, team, contact, department) VALUES (?, ?, ?, ?, ?)",
                           (name, role, team, contact, department))
            conn.commit()
            st.success(f"Team Member '{name}' added successfully!")
        elif submitted:
            st.error("Name and Role are required fields!")

    # Displaying Organizing Team Members
    st.header("Organizing Team Members")
    
    # Fetching team members from the database
    team_members = pd.read_sql_query("SELECT * FROM organizing_team", conn)

    # If team members exist, display in a structured table format
    if not team_members.empty:
        # Display Table with Members
        st.write("""
        <style>
            .dataframe th, .dataframe td {
                padding: 12px;
                text-align: left;
                border: 1px solid #ddd;
                font-size: 14px;
            }
            .delete-btn {
                background-color: red;
                color: white;
                border: none;
                cursor: pointer;
                padding: 6px 12px;
                border-radius: 5px;
            }
            .delete-btn:hover {
                background-color: darkred;
            }
        </style>
        """, unsafe_allow_html=True)

        # Create a table to display the members with delete buttons
        st.table(team_members)

        for _, row in team_members.iterrows():
            delete_button = st.button(f"Delete {row['name']}", key=f"delete_{row['id']}")
            if delete_button:
                cursor.execute("DELETE FROM organizing_team WHERE id = ?", (row['id'],))
                conn.commit()
                st.warning(f"Team Member '{row['name']}' deleted successfully!")
                st.experimental_rerun()  # Refresh the page to reflect changes
    else:
        st.write("No team members found.")


        
    # # Display Tasks
    # st.subheader("Task List")
    # tasks = pd.read_sql_query("SELECT * FROM tasks", conn)
    # for _, row in tasks.iterrows():
    #     st.markdown(f"### {row['task_name']} (Priority: {row['priority']})")
    #     st.write(f"Assigned to: {row['assigned_to']} | Due: {row['due_date']} | Status: {row['status']} | Team: {row['team']}")
    #     col1, col2 = st.columns(2)
    #     with col1:
    #         if row['status'] == "Pending" and st.button("Mark as Completed", key=f"complete_{row['id']}"):
    #             cursor.execute("UPDATE tasks SET status='Completed' WHERE id=?", (row['id'],))
    #             conn.commit()
    #             st.success(f"Task '{row['task_name']}' marked as completed!")
    #     with col2:
    #         if st.button("Delete Task", key=f"delete_{row['id']}"):
    #             cursor.execute("DELETE FROM tasks WHERE id=?", (row['id'],))
    #             conn.commit()
    #             st.warning(f"Task '{row['task_name']}' deleted!")


                # Tasks Management
elif menu == "✅ Tasks Management":
    st.title("Tasks Management")
    
    # Add New Task
    st.subheader("Assign New Task")
    with st.form("add_task_form"):
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            task_name = st.text_input("Task Name")
        with col2:
            assigned_to = st.selectbox(
                "Assign To",
                [f"{row[0]} - {row[1]} ({row[2]})" for row in cursor.execute("SELECT id, name, role FROM organizing_team").fetchall()]
            )
        with col3:
            due_date = st.date_input("Due Date", datetime.now())
        with col4:
            priority = st.selectbox("Priority", ["Low", "Medium", "High"])
        
        # Dropdown for selecting a team
        team = st.selectbox("Team Name", [
            "Event Management Team",
            "Creative Team",
            "Design Team",
            "Marketing Team",
            "Sponsorship Team",
            "Hospitality Team",
            "Media Team",
            "Documentation Team",
            "Tech Team"
        ])
        
        submitted = st.form_submit_button("Assign Task")

        if submitted and task_name and assigned_to and team:
            assignee_name = assigned_to.split(" - ")[1].split(" (")[0]  # Extract name from dropdown text
            cursor.execute("INSERT INTO tasks (task_name, assigned_to, due_date, priority, team) VALUES (?, ?, ?, ?, ?)",
                           (task_name, assignee_name, due_date, priority, team))
            conn.commit()
            st.success(f"Task '{task_name}' assigned to {assignee_name} for team '{team}'!")
        elif submitted:
            st.error("Task Name, Assignee, and Team are required!")

    # Display Tasks as a Table
    st.subheader("Task List")
    
    tasks = pd.read_sql_query("SELECT * FROM tasks", conn)

    # Show tasks in a table format with action buttons
    if not tasks.empty:
        for _, row in tasks.iterrows():
            col1, col2, col3, col4 = st.columns([3, 2, 1, 1])
            with col1:
                st.write(f"**{row['task_name']}** (Priority: {row['priority']})")
            with col2:
                st.write(f"Assigned to: {row['assigned_to']}")
                st.write(f"Due: {row['due_date']}")
            with col3:
                st.write(f"Status: {row['status']}")
                if row['status'] == "Pending":
                    if st.button(f"Mark {row['task_name']} as Completed", key=f"complete_{row['id']}"):
                        cursor.execute("UPDATE tasks SET status='Completed' WHERE id=?", (row['id'],))
                        conn.commit()
                        st.success(f"Task '{row['task_name']}' marked as completed!")
            with col4:
                if st.button(f"Delete {row['task_name']}", key=f"delete_{row['id']}"):
                    cursor.execute("DELETE FROM tasks WHERE id=?", (row['id'],))
                    conn.commit()
                    st.warning(f"Task '{row['task_name']}' deleted!")
    else:
        st.warning("No tasks found.")

elif menu == "🛡️ Team Pages":
    st.title("Team Pages")
    
    # Fetch and list all teams dynamically
    try:
        teams = cursor.execute("SELECT DISTINCT team FROM organizing_team WHERE team != ''").fetchall()
    except Exception as e:
        st.error(f"Error fetching teams: {e}")
        teams = []
    
    selected_team = st.selectbox("Select a Team", [team[0] for team in teams] if teams else [])
    
    if selected_team:
        st.header(f"Team: {selected_team}")
        
        # Use tabs for better navigation
        tab1, tab2 = st.tabs(["Team Members", "Tasks"])

        # Tab 1: Display Team Members
        with tab1:
            st.subheader("Team Members")
            try:
                team_members = pd.read_sql_query(
                    f"SELECT * FROM organizing_team WHERE team = '{selected_team}'", conn
                )
                st.dataframe(team_members)
            except Exception as e:
                st.error(f"Error fetching team members: {e}")

        # Tab 2: Display and Manage Tasks
        with tab2:
            st.subheader("Tasks Assigned to this Team")
            try:
                tasks = pd.read_sql_query(
                    f"SELECT * FROM tasks WHERE team = '{selected_team}' ORDER BY priority DESC", conn
                )
                
                if tasks.empty:
                    st.info("No tasks assigned to this team.")
                else:
                    # Table Header
                    col1, col2, col3, col4, col5, col6 = st.columns([2, 1, 2, 2, 1, 1])
                    col1.write("Task Name")
                    col2.write("Priority")
                    col3.write("Assigned To")
                    col4.write("Due Date")
                    col5.write("Status")
                    col6.write("Complete Task")

                    # Populate table rows dynamically
                    for idx, row in tasks.iterrows():
                        col1, col2, col3, col4, col5, col6 = st.columns([2, 1, 2, 2, 1, 1])
                        col1.write(row["task_name"])
                        col2.write(row["priority"])
                        col3.write(row["assigned_to"])
                        col4.write(row["due_date"])
                        col5.write(row["status"])

                        # Complete Task Button
                        if col6.button("✅", key=f"complete_{row['id']}"):
                            try:
                                cursor.execute("UPDATE tasks SET status=? WHERE id=?", ("Completed", row['id']))
                                conn.commit()
                                st.experimental_rerun()
                            except Exception as e:
                                st.error(f"Error completing task: {e}")

            except Exception as e:
                st.error(f"Error fetching tasks: {e}")


# Export Data (CSV/PDF/DOCX)

if menu == "📂 Export Data":
    st.title("Export Hackathon Data")

    st.markdown("""
    Select a team below to export the data related to their participation, including both the team details and their assigned tasks.
    You can download the data in CSV, PDF, or DOCX format.
    """)

    # Fetch all teams for selection
    teams = cursor.execute("SELECT DISTINCT team FROM organizing_team WHERE team != ''").fetchall()
    selected_team = st.selectbox("Choose a Team", [team[0] for team in teams] if teams else ["No teams available"])

    if selected_team:
        st.subheader(f"Export Data for Team: {selected_team}")

        # Fetch team-specific data
        team_data = pd.read_sql_query(f"SELECT * FROM organizing_team WHERE team = '{selected_team}'", conn)

        # Fetch task-specific data for the selected team
        tasks_data = pd.read_sql_query(f"SELECT * FROM tasks WHERE team = '{selected_team}'", conn)

        # Display options to download team data
        st.markdown("### Download Team Data")

        # Arrange download buttons in columns (side by side)
        col1, col2, col3 = st.columns(3)

        with col1:
            # Download team data as CSV
            st.download_button(
                label="Download Team Data as CSV",
                data=team_data.to_csv(index=False),
                file_name=f"{selected_team}_team_data.csv",
                mime="text/csv"
            )

        with col2:
            # Download team data as PDF
            if st.button("Team Data as PDF"):
                def add_header_footer(canvas, doc):
                    header_path = "header.png"  # path to the header image
                    footer_path = "footer.png"  # path to the footer image
                    canvas.drawImage(header_path, 5, 740, width=550, height=95, preserveAspectRatio=True)
                    canvas.drawImage(footer_path, 40, 40, width=500, height=50, preserveAspectRatio=True)

                def create_pdf(data, file_name, title):
                    pdf = BytesIO()
                    doc = SimpleDocTemplate(pdf, pagesize=A4)

                    # Add custom title style with more space after it to avoid overlap with the header
                    styles = getSampleStyleSheet()
                    title_style = ParagraphStyle(
                        "TitleStyle",
                        parent=styles['Title'],
                        fontSize=0,
                        spaceBefore=110,  # Add space before title to push it down from the header
                        spaceAfter=20,   # Add space after title for separation from table
                        textColor=colors.black
                    )

                    # Create the title and content for the document
                    content = [Paragraph(f"<b>{title}</b>", title_style)]
                    
                    # Prepare table data
                    table_data = [list(data.columns)] + data.values.tolist()
                    table = Table(table_data)

                    # Style the table
                    table.setStyle(TableStyle([
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("SIZE", (0, 0), (-1, -1), 10),
                        ("BOX", (0, 0), (-1, -1), 0.25, colors.black),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.black)
                    ]))

                    # Add the table to content
                    content.append(table)

                    # Build the PDF
                    doc.build(content, onFirstPage=add_header_footer, onLaterPages=add_header_footer)
                    pdf.seek(0)

                    # Download the PDF
                    st.download_button(
                        label="Download",
                        data=pdf,
                        file_name=file_name,
                        mime="application/pdf"
                    )

                # Call create_pdf for Team Data
                create_pdf(team_data, f"{selected_team}_team_data.pdf", f"Team Data for {selected_team}")

        with col3:
            # Download team data as DOCX
            if st.button("Download Team Data as DOCX"):
                doc = Document()
                doc.add_heading(f"Team: {selected_team}", level=1)

                # Add table for team data
                table = doc.add_table(rows=1, cols=len(team_data.columns))
                for i, column_name in enumerate(team_data.columns):
                    table.cell(0, i).text = str(column_name)

                for row in team_data.itertuples(index=False):
                    cells = table.add_row().cells
                    for i, value in enumerate(row):
                        cells[i].text = str(value)

                doc_file = BytesIO()
                doc.save(doc_file)
                doc_file.seek(0)
                st.download_button(
                    label="Download Team Data as DOCX",
                    data=doc_file,
                    file_name=f"{selected_team}_team_data.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        # Display options to download task data
        st.markdown("### Download Task Data")

        # Arrange download buttons in columns (side by side)
        col1, col2, col3 = st.columns(3)

        with col1:
            # Download tasks data as CSV
            st.download_button(
                label="Download Tasks Data as CSV",
                data=tasks_data.to_csv(index=False),
                file_name=f"{selected_team}_tasks_data.csv",
                mime="text/csv"
            )

        with col2:
            # Download tasks data as PDF
            if st.button("Tasks Data as PDF"):
                def add_header_footer(canvas, doc):
                    header_path = "header.png"  # path to the header image
                    footer_path = "footer.png"  # path to the footer image
                    canvas.drawImage(header_path, 5, 740, width=550, height=95, preserveAspectRatio=True)
                    canvas.drawImage(footer_path, 40, 40, width=500, height=50, preserveAspectRatio=True)
                
                def create_pdf(data, file_name, title):
                    pdf = BytesIO()
                    doc = SimpleDocTemplate(pdf, pagesize=A4)

                    # Add custom title style with more space after it to avoid overlap with the header
                    styles = getSampleStyleSheet()
                    title_style = ParagraphStyle(
                        "TitleStyle",
                        parent=styles['Title'],
                        fontSize=0,
                        spaceBefore=20,  # Add space before title to push it down from the header
                        spaceAfter=24,   # Add space after title for separation from table
                        textColor=colors.black
                    )

                    content = [Paragraph(f"<b>{title}</b>", title_style)]

                    # Prepare table data
                    table_data = [list(data.columns)] + data.values.tolist()
                    table = Table(table_data)

                    # Style the table
                    table.setStyle(TableStyle([
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("SIZE", (0, 0), (-1, -1), 10),
                        ("BOX", (0, 0), (-1, -1), 0.25, colors.black),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.black)
                    ]))

                    content.append(table)

                    # Build the PDF
                    doc.build(content, onFirstPage=add_header_footer, onLaterPages=add_header_footer)
                    pdf.seek(0)

                    # Download the PDF
                    st.download_button(
                        label="Download",
                        data=pdf,
                        file_name=file_name,
                        mime="application/pdf"
                    )

                # Call create_pdf for Task Data
                create_pdf(tasks_data, f"{selected_team}_tasks_data.pdf", f"Tasks Data for {selected_team}")

        with col3:
            # Download tasks data as DOCX
            if st.button("Download Tasks Data as DOCX"):
                doc = Document()
                doc.add_heading(f"Tasks for Team: {selected_team}", level=1)

                # Add table for tasks data
                table = doc.add_table(rows=1, cols=len(tasks_data.columns))
                for i, column_name in enumerate(tasks_data.columns):
                    table.cell(0, i).text = str(column_name)

                for row in tasks_data.itertuples(index=False):
                    cells = table.add_row().cells
                    for i, value in enumerate(row):
                        cells[i].text = str(value)

                doc_file = BytesIO()
                doc.save(doc_file)
                doc_file.seek(0)
                st.download_button(
                    label="Download Tasks Data as DOCX",
                    data=doc_file,
                    file_name=f"{selected_team}_tasks_data.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Attendence'
elif menu == "📝 Attendence":
    st.sidebar.markdown("<p style='font-size: 14px; color: #777;'>Get help and support for any queries you may have.</p>", unsafe_allow_html=True)

    # Hackathon Attendance System
    if "attendance_data" not in st.session_state:
        st.session_state["attendance_data"] = pd.DataFrame(
            columns=["Member Name", "Roll No", "Department", "Year", "Team", "Role", "Checked In"]
        )

    # Title
    st.title("Hackathon Attendance System")

    # Section: Member Registration
    st.header("Register Participants")
    with st.form("registration_form"):
        name = st.text_input("Member Name")
        roll_no = st.text_input("Roll No")
        department = st.selectbox("Department", ["CO", "AIML", "DS", "ECS", "Other"])
        year = st.selectbox("Year", ["1st Year", "2nd Year", "3rd Year", "4th Year"])
        team = st.selectbox(
            "Team",
            ["Management", "Creative", "Sponsorship", "Marketing", "Design", "Tech", "Hospitality", "Media", "Documentation"],
        )
        role = st.selectbox("Role", ["Event Head", "Guide", "Team Lead", "Member", "Volunteer"])
        submitted = st.form_submit_button("Register")
        if submitted:
            new_entry = {
                "Member Name": name,
                "Roll No": roll_no,
                "Department": department,
                "Year": year,
                "Team": team,
                "Role": role,
                "Checked In": True,  # Automatically mark attendance
            }
            st.session_state["attendance_data"] = pd.concat(
                [st.session_state["attendance_data"], pd.DataFrame([new_entry])], ignore_index=True
            )
            st.success(f"Registered {name} in Team {team}! Attendance has been automatically marked.")

    # Section: Show Attendance Data
    st.header("Attendance Data")
    if not st.session_state["attendance_data"].empty:
        st.dataframe(st.session_state["attendance_data"])

    # Section: Download Attendance as PDF
    st.header("Download Attendance")
    if st.button("Download Attendance as PDF"):

        class CustomPDF(FPDF):
            def header(self):
                self.image("header.png", x=10, y=5, w=190)
                self.set_y(45)  # Adjust content start position below the header

            def footer(self):
                self.set_y(-20)  # Adjust position of the footer above the bottom
                self.image("footer.png", x=10, y=self.get_y(), w=200)
                self.set_font("Arial", size=8)
                self.cell(0, 10, f"Page {self.page_no()}", align="C")

        pdf = CustomPDF()
        pdf.add_page()

        # Add Title
        pdf.set_font("Arial", style="B", size=14)
        pdf.cell(200, 10, txt="Hackathon Attendance Report", ln=True, align="C")
        pdf.ln(1)

        # Add Date
        current_date = datetime.now().strftime("%Y-%m-%d")
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=f"Date: {current_date}", ln=True, align="L")
        pdf.ln(1)

        # Add Attendance Taken By
        attendance_taken_by = st.text_input("Attendance Taken By", "Shadulla Shaikh - Event Head")
        pdf.cell(200, 10, txt=f"Attendance Taken By: {attendance_taken_by}", ln=True, align="L")
        pdf.ln(1)

        # Define column headers and widths
        headers = ["Sr No", "Member Name", "Roll No", "Department", "Year", "Team", "Role", "Checked In"]
        column_widths = [10, 38, 18, 25, 18, 28, 28, 25]  # Adjust widths as needed

        # Add table headers
        pdf.set_font("Arial", style="B", size=10)

        def add_table_headers():
            for header, width in zip(headers, column_widths):
                pdf.cell(width, 10, header, border=1, align="C")
            pdf.ln()

        # Add rows to the PDF
        pdf.set_font("Arial", size=10)
        max_rows_per_page = 40  # Adjust rows per page as needed
        row_count = 0
        sr_no = 1  # Initialize serial number

        # Loop through the rows and add them to the PDF
        add_table_headers()  # Add headers for the first page

        for index, row in st.session_state["attendance_data"].iterrows():
            # Check if a new page is needed and add page break if needed
            if row_count == max_rows_per_page:
                pdf.add_page()
                add_table_headers()  # Re-add headers for the new page
                row_count = 0

            # Add serial number and row data
            pdf.cell(column_widths[0], 10, str(sr_no), border=1, align="C")  # Add serial number
            for data, width in zip(row.values, column_widths[1:]):  # Skip Sr No width
                pdf.cell(width, 10, str(data), border=1, align="C")
            pdf.ln()

            row_count += 1
            sr_no += 1  # Increment serial number

        # Add Footer with Signature Section
        pdf.ln(10)
        pdf.cell(90, 10, txt="Event Head Signature:", ln=False, align="L")
        pdf.cell(90, 10, txt="Event Faculty Coordinator Signature:", ln=True, align="R")
        pdf.ln(5)
        pdf.cell(90, 10, txt="_____________________________", ln=False, align="L")
        pdf.cell(90, 10, txt="_____________________________", ln=True, align="R")

        # Save and download the PDF
        pdf_file = "attendance_report.pdf"
        pdf.output(pdf_file)
        with open(pdf_file, "rb") as file:
            st.download_button(
                label="Download PDF",
                data=file,
                file_name="attendance_report.pdf",
                mime="application/pdf",
            )

import streamlit as st
from fpdf import FPDF

# Custom PDF class that extends FPDF
class CustomPDF(FPDF):
    def header(self):
        # Add a custom header
        try:
            self.image("header.png", x=10, y=5, w=190)  # Adjust as needed
        except FileNotFoundError:
            self.set_font("Arial", 'B', 12)
            self.cell(0, 10, "Header Image Missing - Add header.png", 0, 1, 'C')
        self.set_y(45)  # Set position for content start below the header

    def footer(self):
        # Add a custom footer
        self.set_y(-20)  # Adjust position for the footer
        try:
            self.image("footer.png", x=10, y=self.get_y(), w=200)  # Adjust as needed
        except FileNotFoundError:
            self.set_font("Arial", size=8)
            self.cell(0, 10, "Footer Image Missing - Add footer.png", 0, 1, 'C')
        self.set_font("Arial", size=8)
        self.cell(0, 10, f"Page {self.page_no()}", 0, 1, 'C')  # Centered page number

# Function to create the document
def create_document(content, notice, alignment, font, signatures, file_name="output.pdf"):
    # Create an instance of CustomPDF
    pdf = CustomPDF()

    # Add a page
    pdf.add_page()

    # Set font
    pdf.set_font(font, size=12)

    # Add a function to increase font size and set font type
    def set_custom_font(pdf, font_size=12, font_type="Times", is_bold=False):
        if is_bold:
            pdf.set_font(font_type, 'B', font_size)  # Bold font
        else:
            pdf.set_font(font_type, '', font_size)  # Regular font

    # Add left and right content
    pdf.set_xy(10, 50)  # Below header

    # Set font size and type for left content
    set_custom_font(pdf, font_size=15, font_type="Times")  # Adjust the font size as needed
    pdf.multi_cell(90, 10, content[0])  # Left content

    # Increase horizontal gap by adjusting the X-coordinate for the right content
    gap = 55  # Increase the gap between left and right content

    # Set font size and type for right content
    set_custom_font(pdf, font_size=14, font_type="Times")  # Adjust the font size as needed
    pdf.set_xy(110 + gap, 50)  # Adjust the X position after left content
    pdf.multi_cell(90, 10, content[1])  # Right content

    # Add notice with reduced gap
    pdf.set_xy(10, 70)  # Reduced the Y position to bring notice closer to the content
    pdf.set_font(font, 'B', 15)  # Bold for notice
    pdf.multi_cell(0, 10, notice, align="C")

    # Add user content, reducing the gap between notice and main content
    pdf.set_xy(10, 88)  # Reduced the Y position to bring content closer to the notice
    pdf.set_font(font, size=12)
    if alignment == "Left":
        pdf.multi_cell(0, 6, content[2], align="L")
    elif alignment == "Center":
        pdf.multi_cell(0, 6, content[2], align="C")
    elif alignment == "Right":
        pdf.multi_cell(0, 6, content[2], align="R")
    elif alignment == "Justify":
        pdf.multi_cell(0, 6, content[2], align="J")

    # Add signature section (centered horizontally on the page)
    pdf.set_xy(10, 200)  # Position signatures below the main content
    pdf.set_font(font, size=12)
    total_signatures = len(signatures)
    space_between_signatures = 60  # Adjust space between signatures

    # Calculate the starting X position for center alignment
    available_width = 190  # Total width of the page minus margins
    total_width_needed = space_between_signatures * total_signatures
    starting_x = (available_width - total_width_needed) / 2 + 10  # Adding left margin

    for i, signature in enumerate(signatures):
        name, designation = signature.split("\n")
        x_position = starting_x + (i * space_between_signatures)  # Adjust spacing between signatures
        pdf.set_xy(x_position, 200)  # Name position
        pdf.cell(50, 10, name, border=0, align='C')  # Name
        pdf.set_xy(x_position, 205)  # Reduced gap between name and designation
        pdf.cell(50, 10, designation, border=0, align='C')  # Designation

    # Output the PDF to a file
    pdf.output(file_name)

# Streamlit interface for the user to input document content and settings
def main():
    st.title("Custom PDF Document Generator with Centered Signatures")

    # Inputs for left and right content
    content_left = st.text_area("Left Content", "Ref: AIKTC/ADMIN/2025/000")
    content_right = st.text_input("Date", "Date: 00/00/0000")
    notice = st.text_input("Notice", "Content Title")
    main_content = st.text_area("Main Content", "This is the main content of the document.")
    content = [content_left, content_right, main_content]

    # Alignment options
    alignment = st.selectbox("Text Alignment", ["Left", "Center", "Right", "Justify"])

    # Font options
    font = st.selectbox("Font", ["Arial", "Times", "Courier", "Helvetica", "Symbol"])

    # Signature inputs
    st.subheader("Signatures")
    num_signatures = st.slider("Number of Signatures", 1, 4, 3)
    signatures = []
    for i in range(num_signatures):
        name = st.text_input(f"Name for Signature {i + 1}", f"Name {i + 1}")
        designation = st.text_input(f"Designation for Signature {i + 1}", f"Designation {i + 1}")
        signatures.append(f"{name}\n{designation}")

    # Button to generate document
    if st.button("Generate Document"):
        # Generate the document
        create_document(content, notice, alignment, font, signatures, "generated_document_with_centered_signatures.pdf")

        # Provide a download button
        with open("generated_document_with_centered_signatures.pdf", "rb") as pdf_file:
            st.download_button(
                label="Download Document",
                data=pdf_file,
                file_name="generated_document_with_centered_signatures.pdf",
                mime="application/pdf"
            )

        st.success("Document generated successfully!")

if __name__ == "__main__":
    main()
# Footer
st.markdown("<footer style='text-align: center; padding: 20px; background-color:Black; color: white;'>© 2024 MammoCare(Silent Echo). All rights reserved.</footer>", unsafe_allow_html=True)
